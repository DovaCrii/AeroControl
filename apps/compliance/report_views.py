"""Compliance status report: on-screen plus CSV / XLSX / DOCX exports.

Follows the Workboard report pattern (same freeze-panes/auto-filter XLSX and
controlled-template DOCX), and reads its numbers from apps.compliance.reports
so the screen, the spreadsheet and the executive email cannot disagree.
"""

import csv
from datetime import datetime, timedelta
from io import BytesIO

from django.core.exceptions import ValidationError
from django.http import HttpResponse
from django.utils import timezone
from django.utils.translation import gettext as _
from django.views.generic import TemplateView, View

from apps.core.exports import neutralize
from apps.core.views import ModelViewPermissionRequiredMixin
from apps.registry.models import CostCenter
from .models import Document, DocumentType
from .reports import (
    ALERT_HEADERS,
    COST_CENTER_HEADERS,
    alert_rows,
    build_compliance_report,
    compare_periods,
    cost_center_rows,
    latest_snapshot_before,
    previous_period,
    totals_from_snapshot,
)

FILENAME_STEM = "aerocontrol-cumplimiento"


def _parse_date(value):
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except (TypeError, ValueError):
        return None


def _lookup_by_pk(queryset, raw_pk):
    """`queryset.filter(pk=raw_pk).first()`, treating a malformed UUID the
    same as "not found" instead of a 500.

    A bookmarked URL, browser autofill, or a bot probing query strings can
    hand this a value that isn't a UUID at all -- `pk` is a UUIDField on both
    CostCenter and DocumentType, and `.filter(pk=...)` raises ValidationError
    (uncaught, a 500) rather than just finding nothing, unlike every other
    mismatch (wrong-but-valid UUID, wrong tenant, archived row).
    """
    try:
        return queryset.filter(pk=raw_pk).first()
    except (ValueError, ValidationError):
        return None


class ComplianceReportMixin(ModelViewPermissionRequiredMixin):
    """Resolves the shared filters, kept in the URL like the other reports."""

    model = Document
    permission_action = "view"

    def _filters(self, request):
        # Same timezone as the report's own period bounds; see reports.py.
        end = _parse_date(request.GET.get("end")) or timezone.localdate()
        start = _parse_date(request.GET.get("start")) or (end - timedelta(days=30))
        cost_center = None
        doc_type = None
        if request.GET.get("cost_center"):
            cost_center = _lookup_by_pk(
                CostCenter.objects.filter(is_active=True), request.GET["cost_center"]
            )
        if request.GET.get("doc_type"):
            doc_type = _lookup_by_pk(
                DocumentType.objects.filter(is_active=True), request.GET["doc_type"]
            )
        return start, end, cost_center, doc_type

    def report_for(self, request):
        start, end, cost_center, doc_type = self._filters(request)
        return build_compliance_report(
            start=start, end=end, cost_center=cost_center, doc_type=doc_type
        )

    def report_and_comparison_for(self, request):
        """R6.4: the current report plus the same current-vs-previous-period
        reading the executive email already sends -- so the web page and the
        inbox tell the same story about the same numbers instead of the web
        page only showing the raw counters.

        R7.7: when a stored snapshot exists from before the period, the
        documentary counters of "previous" come from **it** instead of from a
        recomputation. Without that substitution those three counters are
        always evaluated "as of today" whatever period is asked for, so the
        comparison could only ever read "no change" -- the finding R6.4
        documented and left open. The resolution stats keep coming from the
        recomputed report, because those *do* honour start/end correctly.

        Degrades on purpose: no snapshot (fresh install, or the job never ran)
        means the previous behaviour, not an error.
        """
        start, end, cost_center, doc_type = self._filters(request)
        current = build_compliance_report(
            start=start, end=end, cost_center=cost_center, doc_type=doc_type
        )
        prev_start, prev_end = previous_period(start, end)
        previous = build_compliance_report(
            start=prev_start, end=prev_end, cost_center=cost_center, doc_type=doc_type
        )
        snapshot = None
        # A doc_type filter narrows the report in a way the snapshot does not
        # record (snapshots are unfiltered), so comparing against one would be
        # apples to oranges. Only substitute when the view is unfiltered by type.
        if doc_type is None:
            snapshot = latest_snapshot_before(start, cost_center=cost_center)
        if snapshot is not None:
            previous["totals"] = totals_from_snapshot(snapshot)
        return current, compare_periods(current, previous), snapshot


class ComplianceReportView(ComplianceReportMixin, TemplateView):
    template_name = "compliance/report.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        report, comparison, snapshot = self.report_and_comparison_for(self.request)
        context.update(
            report=report,
            comparison=comparison,
            # R7.7: the template says *where* the comparison baseline came
            # from. "Compared with the previous period" is misleading when
            # there is no history and both sides are the same instant.
            comparison_baseline=snapshot,
            title=_("Compliance status report"),
            cost_centers=CostCenter.objects.filter(is_active=True).order_by("code"),
            document_types=DocumentType.objects.filter(is_active=True).order_by("name"),
            filter_params=self.request.GET,
            export_query=self.request.GET.urlencode(),
        )
        return context


class ComplianceReportCsvView(ComplianceReportMixin, View):
    def get(self, request):
        report = self.report_for(request)
        response = HttpResponse(content_type="text/csv; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{FILENAME_STEM}.csv"'
        response.write("﻿")  # BOM so Excel detects UTF-8
        writer = csv.writer(response, lineterminator="\r\n")
        writer.writerow([_("Compliance status report")])
        writer.writerow([_("Generated"), report["generated_on"].isoformat()])
        writer.writerow([])
        writer.writerow(COST_CENTER_HEADERS)
        for row in cost_center_rows(report):
            writer.writerow([neutralize(value) for value in row])
        totals = report["totals"]
        writer.writerow(
            [
                neutralize(_("TOTAL")),
                "",
                totals["total"],
                totals["valid"],
                totals["valid_pct"],
                totals["expired"],
                totals["due_7"],
                totals["due_15"],
                totals["due_30"],
            ]
        )
        writer.writerow([])
        writer.writerow([_("Open alerts")])
        writer.writerow(ALERT_HEADERS)
        for row in alert_rows(report):
            writer.writerow([neutralize(value) for value in row])
        return response


class ComplianceReportXlsxView(ComplianceReportMixin, View):
    def get(self, request):
        from openpyxl import Workbook
        from openpyxl.styles import Font

        report = self.report_for(request)
        workbook = Workbook()
        summary = workbook.active
        summary.title = "Resumen"
        summary.append(COST_CENTER_HEADERS)
        for row in cost_center_rows(report):
            summary.append([neutralize(value) for value in row])
        totals = report["totals"]
        summary.append(
            [
                _("TOTAL"),
                "",
                totals["total"],
                totals["valid"],
                totals["valid_pct"],
                totals["expired"],
                totals["due_7"],
                totals["due_15"],
                totals["due_30"],
            ]
        )
        for cell in summary[1]:
            cell.font = Font(bold=True)
        for cell in summary[summary.max_row]:
            cell.font = Font(bold=True)

        alerts = workbook.create_sheet("Alertas abiertas")
        alerts.append(ALERT_HEADERS)
        for cell in alerts[1]:
            cell.font = Font(bold=True)
        for row in alert_rows(report):
            alerts.append([neutralize(value) for value in row])

        for sheet in (summary, alerts):
            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions
            for column in sheet.columns:
                sheet.column_dimensions[column[0].column_letter].width = min(
                    max(len(str(cell.value or "")) for cell in column) + 2, 40
                )

        output = BytesIO()
        workbook.save(output)
        response = HttpResponse(
            output.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ),
        )
        response["Content-Disposition"] = f'attachment; filename="{FILENAME_STEM}.xlsx"'
        return response


def build_report_workbook_bytes(report):
    """XLSX bytes for a prepared report, for attaching to the executive email."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Resumen"
    sheet.append(COST_CENTER_HEADERS)
    for cell in sheet[1]:
        cell.font = Font(bold=True)
    for row in cost_center_rows(report):
        sheet.append([neutralize(value) for value in row])
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


class ComplianceReportDocxView(ComplianceReportMixin, View):
    def get(self, request):
        from docx import Document as DocxDocument
        from docx.shared import Inches

        report = self.report_for(request)
        document = DocxDocument()
        document.add_heading(_("AeroControl — Compliance status report"), 0)
        document.add_paragraph(
            _("Generated: %(date)s") % {"date": timezone.localdate().isoformat()}
        )
        period = report["period"]
        document.add_paragraph(
            _("Period analysed: %(start)s to %(end)s")
            % {"start": period["start"].isoformat(), "end": period["end"].isoformat()}
        )
        totals = report["totals"]
        document.add_paragraph(
            _(
                "%(valid)s of %(total)s current documents are valid "
                "(%(pct)s%%). %(expired)s expired, %(due)s expiring within "
                "%(horizon)s days."
            )
            % {
                "valid": totals["valid"],
                "total": totals["total"],
                "pct": totals["valid_pct"],
                "expired": totals["expired"],
                "due": totals["due_7"] + totals["due_15"] + totals["due_30"],
                "horizon": report["horizon_days"],
            }
        )

        document.add_heading(_("By cost center"), level=1)
        table = document.add_table(rows=1, cols=len(COST_CENTER_HEADERS))
        table.style = "Light Shading Accent 1"
        for cell, header in zip(table.rows[0].cells, COST_CENTER_HEADERS):
            cell.text = str(header)
        for row in cost_center_rows(report):
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = str(value)

        resolution = report["resolution"]
        document.add_heading(_("Alert resolution in the period"), level=1)
        if resolution["resolved_count"]:
            document.add_paragraph(
                _("%(count)s alerts resolved, %(days)s days on average.")
                % {
                    "count": resolution["resolved_count"],
                    "days": resolution["avg_days"],
                }
            )
        else:
            document.add_paragraph(_("No alerts were resolved in this period."))

        document.add_heading(_("Open alerts"), level=1)
        if report["open_alerts"]:
            alert_table = document.add_table(rows=1, cols=len(ALERT_HEADERS))
            alert_table.style = "Light Shading Accent 1"
            for cell, header in zip(alert_table.rows[0].cells, ALERT_HEADERS):
                cell.text = str(header)
            for row in alert_rows(report):
                cells = alert_table.add_row().cells
                for cell, value in zip(cells, row):
                    cell.text = str(value)
        else:
            document.add_paragraph(_("No open alerts."))

        for section in document.sections:
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        output = BytesIO()
        document.save(output)
        response = HttpResponse(
            output.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
        response["Content-Disposition"] = f'attachment; filename="{FILENAME_STEM}.docx"'
        return response


class ComplianceReportPdfView(ComplianceReportMixin, View):
    """R6.4: a printable, one-file report -- CSV/XLSX/DOCX all need the
    receiving app installed; a PDF opens the same everywhere and is what
    ISO auditors expect to be handed. Built with reportlab (pure Python,
    no system package like Cairo/Pango or wkhtmltopdf on the Ubuntu VM
    deploy) rather than rendering the HTML template."""

    def get(self, request):
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        report, comparison, _snapshot = self.report_and_comparison_for(request)
        styles = getSampleStyleSheet()
        elements = [
            Paragraph(
                str(_("AeroControl — Compliance status report")), styles["Title"]
            ),
            Paragraph(
                str(_("Generated: %(date)s"))
                % {"date": timezone.localdate().isoformat()},
                styles["Normal"],
            ),
            Paragraph(
                str(_("Period analysed: %(start)s to %(end)s"))
                % {
                    "start": report["period"]["start"].isoformat(),
                    "end": report["period"]["end"].isoformat(),
                },
                styles["Normal"],
            ),
            Spacer(1, 0.2 * inch),
        ]

        def add_table(heading, headers, rows, empty_message):
            elements.append(Paragraph(str(heading), styles["Heading2"]))
            if rows:
                data = [[str(cell) for cell in headers]] + [
                    [str(neutralize(cell)) for cell in row] for row in rows
                ]
                table = Table(data, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1b2a4a")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("FONTSIZE", (0, 0), (-1, -1), 8),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dde3ec")),
                            (
                                "ROWBACKGROUNDS",
                                (0, 1),
                                (-1, -1),
                                [colors.white, colors.HexColor("#f4f6f9")],
                            ),
                        ]
                    )
                )
                elements.append(table)
            else:
                elements.append(Paragraph(str(empty_message), styles["Normal"]))
            elements.append(Spacer(1, 0.25 * inch))

        totals = report["totals"]
        elements.append(
            Paragraph(
                str(
                    _(
                        "%(valid)s of %(total)s current documents are valid "
                        "(%(pct)s%%). %(expired)s expired, %(due)s expiring within "
                        "%(horizon)s days."
                    )
                )
                % {
                    "valid": totals["valid"],
                    "total": totals["total"],
                    "pct": totals["valid_pct"],
                    "expired": totals["expired"],
                    "due": totals["due_7"] + totals["due_15"] + totals["due_30"],
                    "horizon": report["horizon_days"],
                },
                styles["Normal"],
            )
        )
        elements.append(Spacer(1, 0.2 * inch))

        add_table(
            _("Compared with the previous period"),
            [_("KPI"), _("This period"), _("Change"), _("Previous period")],
            [
                [row["label"], row["current"], f"{row['delta']:+g}", row["previous"]]
                for row in comparison
            ],
            _("No comparison available."),
        )
        add_table(
            _("By cost center"),
            COST_CENTER_HEADERS,
            cost_center_rows(report),
            _("No cost centers to report."),
        )

        resolution = report["resolution"]
        elements.append(
            Paragraph(str(_("Alert resolution in the period")), styles["Heading2"])
        )
        if resolution["resolved_count"]:
            elements.append(
                Paragraph(
                    str(_("%(count)s alerts resolved, %(days)s days on average."))
                    % {
                        "count": resolution["resolved_count"],
                        "days": resolution["avg_days"],
                    },
                    styles["Normal"],
                )
            )
        else:
            elements.append(
                Paragraph(
                    str(_("No alerts were resolved in this period.")), styles["Normal"]
                )
            )
        elements.append(Spacer(1, 0.25 * inch))

        add_table(
            _("Open alerts"),
            ALERT_HEADERS,
            alert_rows(report),
            _("No open alerts."),
        )

        output = BytesIO()
        SimpleDocTemplate(
            output,
            pagesize=letter,
            leftMargin=0.6 * inch,
            rightMargin=0.6 * inch,
        ).build(elements)
        response = HttpResponse(output.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{FILENAME_STEM}.pdf"'
        return response
