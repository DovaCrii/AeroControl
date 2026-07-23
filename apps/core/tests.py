import pytest
from datetime import date, timedelta
import json
from docx import Document as DocxDocument
from django.conf import settings
from django.contrib.auth.models import User
from django.contrib.contenttypes.models import ContentType
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.exceptions import ValidationError
from django.contrib.staticfiles import finders
from django.contrib.staticfiles.views import serve
from django.core.management import call_command
from django.core.management.base import CommandError
from django.test import Client, RequestFactory
from django.urls import reverse
from django.utils import timezone

from apps.registry.models import Aircraft, CostCenter, Operator
from apps.compliance.forms import DocumentForm
from apps.compliance.models import Document, DocumentType, document_upload_path
from apps.maintenance.models import MaintenanceRecord
from apps.operations.models import FlightPermission
from apps.workboard.models import KanbanBoard, KanbanStage, KanbanTask
from apps.core.models import AuditEvent, ImportBatch


@pytest.fixture
def client():
    return Client()


@pytest.fixture
def admin_user(db):
    return User.objects.create_superuser("admin", "admin@test.com", "admin123")


@pytest.fixture
def auth_client(client, admin_user):
    assert client.login(username="admin", password="admin123")
    return client


class TestPublicURLs:
    @pytest.mark.django_db
    def test_responses_include_request_id(self, settings):
        settings.DOCUMENTS_ROOT = settings.BASE_DIR
        response = Client().get(reverse("health"), HTTP_X_REQUEST_ID="test-request-42")
        assert response.status_code == 200
        assert response.headers["X-Request-ID"] == "test-request-42"
        assert "Content-Security-Policy-Report-Only" in response.headers

    @pytest.mark.django_db
    def test_health_endpoint_reports_dependencies(self, settings):
        settings.DOCUMENTS_ROOT = settings.BASE_DIR
        response = Client().get(reverse("health"))
        assert response.status_code == 200
        payload = response.json()
        assert payload["status"] == "ok"
        assert payload["checks"]["database"] == "ok"

    @pytest.mark.django_db
    def test_unified_calendar_events_combines_operations_maintenance_and_tasks(self, auth_client):
        center = CostCenter.objects.create(code="CAL", name="Calendar")
        operator = Operator.objects.create(employee_id="CAL-1", full_name="Calendar Operator", cost_center=center)
        aircraft = Aircraft.objects.create(registration="CC-CAL", type="Fixed", model="C1", manufacturer="Maker", cost_center=center)
        permission = FlightPermission.objects.create(
            permission_number="CAL-001", operator=operator, aircraft=aircraft,
            cost_center=center, purpose="Inspection", flight_date=date(2026, 7, 24), location="Santiago",
        )
        maintenance = MaintenanceRecord.objects.create(
            aircraft=aircraft, maintenance_type="scheduled", description="Inspection",
            scheduled_date=date(2026, 7, 25), performed_by="Team", cost=0,
        )
        board = KanbanBoard.objects.create(name="Calendar board")
        stage = KanbanStage.objects.create(board=board, name="Planned")
        task = KanbanTask.objects.create(board=board, stage=stage, title="Calendar task", due_date=date(2026, 7, 26))

        response = auth_client.get(reverse("calendar-events"), {"start": "2026-07-01", "end": "2026-08-01"})

        assert response.status_code == 200
        payload = response.json()
        assert {item["type"] for item in payload} == {"permission", "maintenance", "task"}
        assert {item["id"] for item in payload} == {f"permission-{permission.pk}", f"maintenance-{maintenance.pk}", f"task-{task.pk}"}

    def test_administration_center_explains_operational_configuration(self, auth_client):
        response = auth_client.get(reverse("administration"))

        assert response.status_code == 200
        content = response.content.decode()
        assert "Centro de administración" in content
        assert "Configuración del tablero de trabajo" in content
        assert "Administración técnica avanzada" in content

    """Verify pages that are intentionally available without authentication."""

    def test_login_page(self, client):
        response = client.get(reverse("login"))

        assert response.status_code == 200
        content = response.content.decode()
        assert "Iniciar sesión" in content
        assert "AeroControl" in content

    def test_login_post_ok(self, client, admin_user):
        response = client.post(
            reverse("login"),
            {"username": "admin", "password": "admin123"},
        )

        assert response.status_code == 302
        assert response.url == "/"

    def test_login_post_invalid(self, client, db):
        response = client.post(
            reverse("login"),
            {"username": "unknown", "password": "wrong"},
        )

        assert response.status_code == 200
        content = response.content.decode()
        assert "Usuario o contraseña no válidos" in content


class TestAuthRequiredURLs:
    """Verify that internal pages require authentication and render for users."""

    AUTH_REQUIRED_URLS = [
        "dashboard",
        "alert-count",
        "costcenter-list",
        "aircraft-list",
        "operator-list",
        "assignment-list",
        "document-list",
        "alert-list",
        "permission-list",
        "record-list",
        "calendar",
        "maintenance-list",
        "task-list",
    ]

    @pytest.mark.parametrize("url_name", AUTH_REQUIRED_URLS)
    def test_requires_login(self, client, url_name):
        response = client.get(reverse(url_name))

        assert response.status_code in (302, 403), (
            f"{url_name} should require login, got {response.status_code}"
        )

    @pytest.mark.parametrize("url_name", AUTH_REQUIRED_URLS)
    def test_returns_success_when_authenticated(self, auth_client, url_name):
        response = auth_client.get(reverse(url_name))

        assert response.status_code in (200, 302), (
            f"{url_name} failed with {response.status_code}"
        )


class TestAuthenticatedPages:
    def test_cost_center_import_preview_apply_and_revert(self, auth_client):
        csv_file = SimpleUploadedFile("centers.csv", b"code,name\nIMP-1,Imported one\nIMP-2,Imported two\n", content_type="text/csv")
        preview = auth_client.post(reverse("costcenter-import"), {"file": csv_file})
        assert preview.status_code == 200
        assert "2 rows are valid" in preview.content.decode()

        csv_file = SimpleUploadedFile("centers.csv", b"code,name\nIMP-1,Imported one\nIMP-2,Imported two\n", content_type="text/csv")
        applied = auth_client.post(reverse("costcenter-import"), {"file": csv_file, "apply": "1"})
        assert applied.status_code == 302
        batch = ImportBatch.objects.get(entity="registry.costcenter")
        assert CostCenter.objects.filter(code="IMP-1", is_active=True).exists()

        reverted = auth_client.post(reverse("costcenter-import-revert", args=[batch.pk]))
        assert reverted.status_code == 204
        assert not CostCenter.objects.filter(code="IMP-1", is_active=True).exists()

    def test_aircraft_import_validates_cost_center(self, auth_client):
        center = CostCenter.objects.create(code="AIR-OPS", name="Air Ops")
        payload = b"registration,type,model,manufacturer,year,cost_center,status\nCC-IMP,Drone,X1,Maker,2026,AIR-OPS,active\n"
        response = auth_client.post(
            reverse("aircraft-import"),
            {"file": SimpleUploadedFile("aircraft.csv", payload, content_type="text/csv"), "apply": "1"},
        )
        assert response.status_code == 302
        assert Aircraft.objects.filter(registration="CC-IMP", cost_center=center).exists()

    def test_operator_import_validates_cost_center(self, auth_client):
        center = CostCenter.objects.create(code="OP-OPS", name="Operator Ops")
        payload = b"employee_id,full_name,email,phone,cost_center\nEMP-IMP,Imported Operator,imp@example.com,+56900000000,OP-OPS\n"
        response = auth_client.post(
            reverse("operator-import"),
            {"file": SimpleUploadedFile("operators.csv", payload, content_type="text/csv"), "apply": "1"},
        )
        assert response.status_code == 302
        assert Operator.objects.filter(employee_id="EMP-IMP", cost_center=center).exists()


class TestChapter1DocxImport:
    @pytest.mark.django_db
    def test_docx_source_reports_aircraft_and_duplicate_operators(self, tmp_path, capsys):
        document = DocxDocument()
        document.add_paragraph("1.5- DOTACIÓN OPERADOR RPA")
        document.add_paragraph("1) PERMANENTES")
        document.add_paragraph("1.- NOMBRE : Ana Pérez")
        document.add_paragraph("RUT : 11.111.111-1")
        document.add_paragraph("Credencial N° : 100")
        document.add_paragraph("Tipo : Operador RPA")
        document.add_paragraph("Habilitaciones : Matrice")
        document.add_paragraph("Dirección : Calle 1")
        document.add_paragraph("Teléfono : 999")
        document.add_paragraph("Email : ana@example.com")
        document.add_paragraph("2.- NOMBRE : Ana Pérez")
        document.add_paragraph("RUT : 11.111.111-1")
        document.add_paragraph("Credencial N° : 100")
        document.add_paragraph("Tipo : Operador RPA")
        document.add_paragraph("Habilitaciones : Matrice")
        document.add_paragraph("Dirección : Calle 1")
        document.add_paragraph("Teléfono : 999")
        document.add_paragraph("Email : ana@example.com")
        document.add_paragraph("2) EVENTUALES (NO APLICA)")
        service_table = document.add_table(rows=2, cols=3)
        service_table.rows[0].cells[0].text = "AERONAVES"
        service_table.rows[0].cells[1].text = "REGISTRO"
        service_table.rows[0].cells[2].text = "SERVICIO"
        service_table.rows[1].cells[0].text = "DJI / MAVIC"
        service_table.rows[1].cells[1].text = "RPA 1"
        service_table.rows[1].cells[2].text = "Fotografía"
        inventory = document.add_table(rows=2, cols=8)
        for index, header in enumerate(["Propietario", "Modelo y Año", "Serie", "Inscripción", "MTOW", "Básico", "VLOS", "Paracaídas"]):
            inventory.rows[0].cells[index].text = header
        for index, value in enumerate(["J.E.J.", "DJI / MAVIC", "SER-1", "RPA-1", "1,0", "1,0", "VLOS", "NO"]):
            inventory.rows[1].cells[index].text = value
        source = tmp_path / "chapter1.docx"
        document.save(source)

        call_command("chapter1_docx_import", "--source", str(source), "--json")
        payload = json.loads(capsys.readouterr().out)
        assert payload["counts"] == {
            "aircraft_extracted": 1,
            "operators_extracted": 2,
            "operators_ready": 1,
            "duplicate_groups": 1,
            "cost_centers": 0,
        }

    def test_import_templates_are_downloadable(self, auth_client):
        response = auth_client.get(reverse("aircraft-import"), {"template": "1"})
        assert response.status_code == 200
        assert "registration,type,model" in response.content.decode()

    def test_global_search_respects_permissions(self, auth_client):
        CostCenter.objects.create(code="SEARCH", name="Search Operations")
        response = auth_client.get(reverse("global-search"), {"q": "SEARCH"})
        assert response.status_code == 200
        assert "Search Operations" in response.content.decode()

    def test_global_search_requires_authentication(self, client):
        response = client.get(reverse("global-search"), {"q": "SEARCH"})
        assert response.status_code == 302

    def test_mutating_action_is_audited_without_request_data(self, auth_client):
        response = auth_client.post(reverse("board-create"), {"name": "Audited board"})
        assert response.status_code == 302
        event = AuditEvent.objects.latest("created_at")
        assert event.actor.username == "admin"
        assert event.action == "post_success"
        assert event.path == reverse("board-create")
        assert event.request_id
        assert event.model_label == "workboard.KanbanBoard"
        assert event.object_id
        assert event.metadata == {"query_keys": []}

    def test_audit_write_failure_does_not_change_mutation_response(self, auth_client, monkeypatch):
        def fail_create(**kwargs):
            raise RuntimeError("audit database unavailable")

        monkeypatch.setattr(AuditEvent.objects, "create", fail_create)
        response = auth_client.post(reverse("board-create"), {"name": "Audit resilient board"})
        assert response.status_code == 302
        assert KanbanBoard.objects.filter(name="Audit resilient board").exists()

    def test_audit_events_reject_orm_mutation_and_deletion(self, admin_user):
        event = AuditEvent.objects.create(
            actor=admin_user,
            action="post_success",
            method="POST",
            path="/audit-test/",
            status_code=201,
        )

        event.action = "tampered"
        with pytest.raises(ValidationError):
            event.save()
        with pytest.raises(ValidationError):
            event.delete()
        with pytest.raises(ValidationError):
            AuditEvent.objects.filter(pk=event.pk).update(action="tampered")
        with pytest.raises(ValidationError):
            AuditEvent.objects.filter(pk=event.pk).delete()

    """Verify representative authenticated pages and shared template behavior."""

    def test_dashboard(self, auth_client):
        response = auth_client.get(reverse("dashboard"))

        assert response.status_code == 200
        assert "Dashboard" in response.content.decode()

    def test_base_template_has_htmx(self, auth_client):
        response = auth_client.get(reverse("dashboard"))

        assert response.status_code == 200
        assert "htmx.org" in response.content.decode()

    def test_base_template_has_dark_mode_toggle(self, auth_client):
        response = auth_client.get(reverse("dashboard"))
        content = response.content.decode()

        assert response.status_code == 200
        assert "toggleTheme" in content
        assert "data-theme" in content
        assert "AeroControl" in content

    def test_base_template_has_accessible_navigation_and_modal_hooks(self, auth_client):
        response = auth_client.get(reverse("dashboard"))
        content = response.content.decode()

        assert 'href="#main-content"' in content
        assert 'id="sidebar-toggle"' in content
        assert 'aria-controls="sidebar"' in content
        assert 'aria-labelledby="generic-modal-title"' in content
        assert "shown.bs.modal" in content

    def test_dashboard_serializes_chart_data_without_marking_it_safe(self, auth_client):
        response = auth_client.get(reverse("dashboard"))
        content = response.content.decode()

        assert 'id="chart-data" type="application/json"' in content
        assert "chart_data|safe" not in content

    def test_logout_requires_post_and_logs_out(self, auth_client):
        response = auth_client.post(reverse("logout"))

        assert response.status_code == 302
        assert response.url == reverse("login")

    def test_language_switches_navigation_and_status_labels_to_spanish(self, auth_client):
        response = auth_client.post(
            reverse("set_language"), {"language": "es", "next": reverse("dashboard")}
        )

        assert response.status_code == 302
        response = auth_client.get(reverse("dashboard"))
        assert "Panel de operaciones" in response.content.decode()
        assert "Aeronaves" in response.content.decode()

    def test_language_switches_shared_registry_list_to_spanish(self, auth_client):
        auth_client.post(
            reverse("set_language"), {"language": "es", "next": reverse("dashboard")}
        )

        response = auth_client.get(reverse("costcenter-list"))
        content = response.content.decode()

        assert "Centros de costo" in content
        assert "Exportar CSV" in content
        assert "Todos los estados" in content
        assert "+ Nuevo" in content

    def test_shared_registry_form_labels_are_translated_to_spanish(self, auth_client):
        auth_client.post(
            reverse("set_language"), {"language": "es", "next": reverse("dashboard")}
        )

        response = auth_client.get(
            reverse("costcenter-create"), HTTP_HX_REQUEST="true"
        )
        content = response.content.decode()

        assert response.status_code == 200
        assert "Código" in content
        assert "Nombre" in content

    def test_document_form_labels_and_entity_options_are_translated_to_spanish(
        self, auth_client
    ):
        auth_client.post(
            reverse("set_language"), {"language": "es", "next": reverse("dashboard")}
        )

        response = auth_client.get(
            reverse("document-create"), HTTP_HX_REQUEST="true"
        )
        content = response.content.decode()

        assert response.status_code == 200
        assert "Título" in content
        assert "Tipo de entidad" in content
        assert "Registro de mantenimiento" in content

    @pytest.mark.parametrize(
        ("url_name", "expected"),
        [
            ("alert-list", "Alertas"),
            ("maintenance-list", "Mantenimiento"),
            ("calendar", "Calendario"),
            ("kanban", "Tablero Kanban"),
        ],
    )
    def test_specific_modules_keep_spanish_labels(self, auth_client, url_name, expected):
        auth_client.post(
            reverse("set_language"), {"language": "es", "next": reverse("dashboard")}
        )

        response = auth_client.get(reverse(url_name))

        assert response.status_code == 200
        assert expected in response.content.decode()

    def test_document_create_modal_renders_for_a_new_document(self, auth_client):
        response = auth_client.get(
            reverse("document-create"), HTTP_HX_REQUEST="true"
        )

        assert response.status_code == 200
        assert 'name="title"' in response.content.decode()

    def test_kanban_board_create_redirects_to_board_list(self, auth_client):
        response = auth_client.post(
            reverse("board-create"),
            {"name": "Operations board", "description": "Daily work"},
        )

        assert response.status_code == 302
        assert response.url == reverse("board-list")
        assert KanbanBoard.objects.filter(name="Operations board").exists()

    def test_document_entity_options_limits_selection_to_allowed_records(self, auth_client):
        aircraft_type = ContentType.objects.get_for_model(Aircraft)
        response = auth_client.get(
            reverse("document-entity-options"),
            {"entity_type": aircraft_type.pk},
        )

        assert response.status_code == 200
        assert 'id="document-object-field"' in response.content.decode()

        response = auth_client.get(
            reverse("document-entity-options"), {"entity_type": "invalid"}
        )
        assert response.status_code == 400

    def test_document_upload_rejects_signature_mismatch_and_sanitizes_path(self, db):
        cost_center = CostCenter.objects.create(code="SEC", name="Security")
        aircraft = Aircraft.objects.create(
            registration="CC-SEC",
            type="Fixed",
            model="S",
            manufacturer="Maker",
            cost_center=cost_center,
        )
        doc_type = DocumentType.objects.create(code="certificates", name="Certificates")
        aircraft_type = ContentType.objects.get_for_model(Aircraft)
        form = DocumentForm(
            data={
                "title": "Certificate",
                "doc_type": doc_type.pk,
                "entity_type": aircraft_type.pk,
                "object_id": aircraft.pk,
                "issue_date": date(2026, 1, 1),
                "expiry_date": date(2027, 1, 1),
            },
            files={
                "file": SimpleUploadedFile(
                    "../../evil.pdf", b"not a pdf", content_type="application/pdf"
                )
            },
        )

        assert not form.is_valid()
        assert "file" in form.errors

        document = Document(
            doc_type=doc_type,
            content_type=aircraft_type,
            object_id=aircraft.pk,
        )
        safe_path = document_upload_path(document, "../../evil.pdf")
        assert ".." not in safe_path
        assert safe_path.startswith("certificates/aircraft/")

    def test_document_delete_archives_without_physical_delete(self, auth_client):
        doc_type = DocumentType.objects.create(code="archive-test", name="Archive test")
        aircraft_type = ContentType.objects.get_for_model(Aircraft)
        document = Document.objects.create(
            title="Archive document",
            doc_type=doc_type,
            content_type=aircraft_type,
            object_id="00000000-0000-0000-0000-000000000001",
            file_path="archive-test/aircraft/file.pdf",
            issue_date=date(2026, 1, 1),
        )
        response = auth_client.post(reverse("document-delete", args=[document.pk]))
        assert response.status_code == 302
        document.refresh_from_db()
        assert document.is_active is False


class TestStaticFiles:
    @pytest.mark.django_db
    def test_scale_readiness_reports_database_vendor(self, capsys):
        call_command("scale_readiness")
        output = capsys.readouterr().out
        assert "database_vendor:" in output
        assert "rollback_required: True" in output

    @pytest.mark.django_db
    def test_validate_operational_data_reports_clean_dataset(self, capsys):
        call_command("validate_operational_data")
        output = capsys.readouterr().out
        assert "status: ok" in output
        assert "errors: 0" in output

    @pytest.mark.django_db
    def test_chapter1_mapping_is_versioned(self, capsys):
        call_command("chapter1_mapping", "--json")
        payload = json.loads(capsys.readouterr().out)
        assert payload["version"] == "chapter1-v1"
        assert "aircraft" in payload["entities"]

    def test_chapter1_import_requires_source_files(self):
        with pytest.raises(CommandError, match="source"):
            call_command("chapter1_import")

    def test_css_is_served(self):
        static_path = finders.find("css/app.css")

        assert static_path is not None
        request = RequestFactory().get(f"{settings.STATIC_URL.rstrip('/')}/css/app.css")
        response = serve(request, "css/app.css", insecure=True)

        assert response.status_code == 200
        assert response["Content-Type"].startswith("text/css")

    def test_backup_writes_manifest_and_detects_tampering(self, monkeypatch, tmp_path):
        monkeypatch.setenv("BACKUPS_DIR", str(tmp_path))
        source = tmp_path / "source.sqlite3"
        source.write_bytes(b"sqlite test database")
        monkeypatch.setitem(settings.DATABASES["default"], "NAME", str(source))

        call_command("backup")
        backup = next(tmp_path.glob("*.sqlite3"))
        manifest = backup.with_suffix(".json")

        assert manifest.is_file()
        call_command("verify_backup", str(backup))

        with backup.open("ab") as stream:
            stream.write(b"tampered")
        with pytest.raises(CommandError, match="(size|checksum)"):
            call_command("verify_backup", str(backup))

    def test_backup_can_be_restored_to_explicit_destination(self, monkeypatch, tmp_path):
        monkeypatch.setenv("BACKUPS_DIR", str(tmp_path))
        source = tmp_path / "source.sqlite3"
        source.write_bytes(b"restore me")
        monkeypatch.setitem(settings.DATABASES["default"], "NAME", str(source))

        call_command("backup")
        backup = next(tmp_path.glob("*.sqlite3"))
        destination = tmp_path / "restored" / "db.sqlite3"
        call_command("restore_backup", str(backup), str(destination))

        assert destination.read_bytes() == b"restore me"

    def test_cleanup_documents_dry_run_and_execute(self, monkeypatch, db, tmp_path):
        monkeypatch.setattr(settings, "DOCUMENTS_ROOT", tmp_path)
        doc_type = DocumentType.objects.create(code="retention", name="Retention")
        aircraft_type = ContentType.objects.get_for_model(Aircraft)
        document = Document.objects.create(
            title="Old certificate",
            doc_type=doc_type,
            content_type=aircraft_type,
            object_id="00000000-0000-0000-0000-000000000001",
            file_path="retention/aircraft/old.pdf",
            issue_date=date(2020, 1, 1),
            expiry_date=None,
        )
        Document.objects.filter(pk=document.pk).update(
            is_active=False, updated_at=timezone.now() - timedelta(days=4000)
        )
        stored = tmp_path / document.file_path
        stored.parent.mkdir(parents=True)
        stored.write_bytes(b"old")

        call_command("cleanup_documents", "--older-than-days", "3650")
        assert stored.exists()
        call_command("cleanup_documents", "--older-than-days", "3650", "--execute")

        document.refresh_from_db()
        assert not stored.exists()
        assert document.file_path == ""


@pytest.mark.parametrize(
    "url_name",
    [
        "costcenter-list",
        "document-list",
        "alert-list",
        "maintenance-list",
        "permission-list",
        "record-list",
        "task-list",
    ],
)
def test_unprivileged_user_cannot_read_or_export_lists(client, db, url_name):
    User.objects.create_user("viewer", password="password")
    assert client.login(username="viewer", password="password")

    response = client.get(reverse(url_name))
    export_response = client.get(reverse(url_name), {"export": "csv"})

    assert response.status_code == 403
    assert export_response.status_code == 403


@pytest.mark.parametrize(
    "url_name",
    [
        "costcenter-create",
        "document-create",
        "maintenance-create",
        "permission-create",
        "task-create",
    ],
)
def test_unprivileged_user_cannot_open_mutating_forms(client, db, url_name):
    User.objects.create_user("viewer", password="password")
    assert client.login(username="viewer", password="password")

    response = client.get(reverse(url_name))

    assert response.status_code == 403
