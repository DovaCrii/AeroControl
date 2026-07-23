import csv
import hashlib
import json
import re
from collections import defaultdict
from decimal import Decimal, InvalidOperation
from pathlib import Path

from django.core.management.base import BaseCommand, CommandError
from django.db import transaction
from docx import Document

from apps.core.models import ImportBatch
from apps.registry.models import Aircraft, CostCenter, Operator


FIELD_LABELS = re.compile(
    r"(?P<label>NOMBRE|RUT|Credencia(?:l)?\s*N[°ºo]?|Tipo|Habilitaciones|"
    r"Direcci[^\s:]*n|Tel[^\s:]*fono|Email)\s*:\s*",
    re.IGNORECASE,
)
RECORD_START = re.compile(
    r"^\s*\d{1,3}\s*[-.]?\s*[-.]?\s*NOMBRE\b",
    re.IGNORECASE | re.MULTILINE,
)


def clean_text(value):
    return re.sub(r"\s+", " ", str(value or "")).strip(" \t:;.")


def rut_key(value):
    return re.sub(r"[^0-9K]", "", value.upper())


def parse_decimal(value):
    value = clean_text(value).replace(",", ".")
    if not value:
        return None
    try:
        return Decimal(value)
    except InvalidOperation as exc:
        raise CommandError(f"Invalid weight value: {value}") from exc


class Command(BaseCommand):
    help = "Validate and optionally apply the official Chapter 1 DOCX source."

    def add_arguments(self, parser):
        parser.add_argument("--source", type=Path, required=True)
        parser.add_argument("--cost-centers", type=Path)
        parser.add_argument("--export-dir", type=Path)
        parser.add_argument("--apply", action="store_true")
        parser.add_argument("--json", action="store_true", dest="as_json")

    def read_cost_centers(self, path):
        if not path:
            return []
        if not path.is_file():
            raise CommandError(f"Missing cost center source: {path}")
        with path.open("r", encoding="utf-8-sig", newline="") as stream:
            reader = csv.DictReader(stream)
            expected = ["code", "name", "responsible"]
            if reader.fieldnames != expected:
                raise CommandError(f"cost centers columns must be: {','.join(expected)}")
            rows = [
                {field: clean_text(value) for field, value in row.items()}
                for row in reader
            ]
        codes = [row["code"] for row in rows]
        if any(not code for code in codes) or len(codes) != len(set(codes)):
            raise CommandError("cost centers contains empty or duplicate codes")
        if any(not row["name"] for row in rows):
            raise CommandError("cost centers requires a name for every code")
        return rows

    def extract_operators(self, document):
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        start_match = re.search(r"1\.5\s*[-.]?\s*DOTACI", text, re.IGNORECASE)
        end_match = re.search(r"EVENTUALES\s*\(\s*NO APLICA", text, re.IGNORECASE)
        if not start_match:
            raise CommandError("Could not locate the permanent operator section")
        section = text[start_match.start() : end_match.start() if end_match else None]
        starts = [match.start() for match in RECORD_START.finditer(section)]
        records = []
        for index, boundary in enumerate(starts):
            block = section[boundary : starts[index + 1] if index + 1 < len(starts) else None]
            values = {
                "full_name": "",
                "rut": "",
                "dgac_credential": "",
                "operator_type": "",
                "authorizations": "",
                "address": "",
                "phone": "",
                "email": "",
            }
            matches = list(FIELD_LABELS.finditer(block))
            for position, match in enumerate(matches):
                raw_label = re.sub(r"[^a-z]", "", match.group("label").lower())
                field = {
                    "nombre": "full_name",
                    "rut": "rut",
                    "credencialn": "dgac_credential",
                    "credencian": "dgac_credential",
                    "tipo": "operator_type",
                    "habilitaciones": "authorizations",
                    "direccion": "address",
                    "direccin": "address",
                    "telefono": "phone",
                    "telfono": "phone",
                    "email": "email",
                }.get(raw_label)
                if not field:
                    continue
                end = matches[position + 1].start() if position + 1 < len(matches) else len(block)
                values[field] = clean_text(block[match.end() : end])
                if field == "email":
                    values[field] = values[field].split("2)", 1)[0].strip()
            if not values["full_name"] or not values["rut"]:
                continue
            values["source_index"] = len(records) + 1
            records.append(values)
        if not records:
            raise CommandError("No permanent operators were extracted from the source")
        return records

    def extract_aircraft(self, document):
        if len(document.tables) < 2:
            raise CommandError("The source must contain the aircraft inventory table")
        service_table = document.tables[0]
        shared_services = clean_text(service_table.rows[1].cells[2].text if len(service_table.rows) > 1 else "")
        inventory = document.tables[1]
        aircraft = []
        for row in inventory.rows[1:]:
            values = [clean_text(cell.text) for cell in row.cells]
            if len(values) < 8 or not values[3]:
                continue
            model_text = values[1]
            manufacturer, model = (model_text.split("/", 1) + [""])[:2]
            if not model:
                manufacturer, model = "", model_text
            registration_match = re.search(r"RPA\s*[- ]?\s*(\d+)", values[3], re.IGNORECASE)
            registration = f"RPA-{registration_match.group(1)}" if registration_match else values[3]
            aircraft.append(
                {
                    "registration": registration,
                    "type": "RPA",
                    "model": clean_text(model),
                    "manufacturer": clean_text(manufacturer),
                    "year": None,
                    "serial_number": values[2],
                    "max_takeoff_weight_kg": parse_decimal(values[4]),
                    "basic_weight_kg": parse_decimal(values[5]),
                    "vlos": values[6],
                    "parachute": values[7],
                    "authorized_services": shared_services,
                }
            )
        if not aircraft:
            raise CommandError("No aircraft were extracted from the source")
        return aircraft

    def duplicate_report(self, operators):
        groups = defaultdict(list)
        for operator in operators:
            groups[rut_key(operator["rut"])].append(operator)
        duplicate_groups = []
        clean_records = []
        for key, records in groups.items():
            if len(records) == 1:
                clean_records.append(records[0])
                continue
            comparable = [
                tuple(record[field] for field in ("full_name", "email", "phone", "address", "authorizations"))
                for record in records
            ]
            kind = "exact_duplicate" if len(set(comparable)) == 1 else "conflicting_duplicate"
            duplicate_groups.append({"rut": key, "kind": kind, "records": records})
            if kind == "exact_duplicate":
                clean_records.append(records[0])
        return clean_records, duplicate_groups

    def build_report(self, source, cost_centers):
        document = Document(source)
        operators = self.extract_operators(document)
        aircraft = self.extract_aircraft(document)
        clean_operators, duplicate_groups = self.duplicate_report(operators)
        digest = hashlib.sha256(source.read_bytes()).hexdigest()
        return {
            "mapping": "chapter1-v1-docx",
            "source": source.name,
            "sha256": digest,
            "cost_centers": cost_centers,
            "aircraft": aircraft,
            "operators": clean_operators,
            "duplicate_groups": duplicate_groups,
            "counts": {
                "aircraft_extracted": len(aircraft),
                "operators_extracted": len(operators),
                "operators_ready": len(clean_operators),
                "duplicate_groups": len(duplicate_groups),
                "cost_centers": len(cost_centers),
            },
        }

    def export_report(self, report, directory):
        directory.mkdir(parents=True, exist_ok=True)
        (directory / "chapter1-report.json").write_text(
            json.dumps(report, ensure_ascii=False, indent=2, default=str), encoding="utf-8"
        )
        for name, rows in (("chapter1-aircraft.csv", report["aircraft"]), ("chapter1-operators.csv", report["operators"])):
            if not rows:
                continue
            fields = [field for field in rows[0] if field != "source_index"]
            with (directory / name).open("w", encoding="utf-8", newline="") as stream:
                writer = csv.DictWriter(stream, fieldnames=fields)
                writer.writeheader()
                writer.writerows({field: row.get(field, "") for field in fields} for row in rows)

    def apply_report(self, report):
        existing_aircraft = set(Aircraft.objects.values_list("registration", flat=True))
        existing_operators = set(Operator.objects.values_list("employee_id", flat=True))
        existing_centers = set(CostCenter.objects.values_list("code", flat=True))
        collisions = [
            *(f"aircraft:{row['registration']}" for row in report["aircraft"] if row["registration"] in existing_aircraft),
            *(f"operator:RUT-{rut_key(row['rut'])}" for row in report["operators"] if f"RUT-{rut_key(row['rut'])}" in existing_operators),
            *(f"cost_center:{row['code']}" for row in report["cost_centers"] if row["code"] in existing_centers),
        ]
        if collisions:
            raise CommandError("Existing records would be overwritten: " + ", ".join(collisions))
        created_ids = []
        with transaction.atomic():
            for row in report["cost_centers"]:
                created_ids.append(str(CostCenter.objects.create(**row).pk))
            source_note = f"Fuente oficial Capítulo 1: {report['source']} ({report['sha256'][:12]})"
            for row in report["aircraft"]:
                payload = {**row, "notes": source_note}
                created_ids.append(str(Aircraft.objects.create(**payload).pk))
            for row in report["operators"]:
                operator_data = dict(row)
                rut = operator_data["rut"]
                payload = {
                    **operator_data,
                    "employee_id": f"RUT-{rut_key(rut)}",
                    "notes": source_note,
                }
                payload.pop("source_index", None)
                created_ids.append(str(Operator.objects.create(**payload).pk))
            stored_report = json.loads(json.dumps(report, ensure_ascii=False, default=str))
            batch = ImportBatch.objects.create(
                actor=None,
                entity="chapter1.docx",
                rows=stored_report,
                created_ids=created_ids,
            )
        return batch

    def handle(self, *args, **options):
        source = options["source"]
        if not source.is_file():
            raise CommandError(f"Missing source DOCX: {source}")
        report = self.build_report(source, self.read_cost_centers(options.get("cost_centers")))
        if options.get("export_dir"):
            self.export_report(report, options["export_dir"])
        if options["apply"]:
            self.apply_report(report)
        output = report | {"apply": options["apply"]}
        if options["as_json"]:
            self.stdout.write(json.dumps(output, ensure_ascii=False, default=str))
            return
        self.stdout.write(f"mapping: {report['mapping']}")
        self.stdout.write(f"source: {report['source']}")
        for key, value in report["counts"].items():
            self.stdout.write(f"{key}: {value}")
        for duplicate in report["duplicate_groups"]:
            self.stdout.write(f"duplicate {duplicate['kind']} RUT {duplicate['rut']}")
