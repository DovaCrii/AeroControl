from django.contrib.auth.mixins import LoginRequiredMixin
import csv
import json
from django.core.exceptions import ValidationError
from django.db import transaction
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils.translation import gettext as _
from django.views import View
from django.views.generic import ListView, CreateView, TemplateView, UpdateView
from uuid import UUID

from apps.core.views import (
    CsvExportMixin,
    ModelPermissionRequiredMixin,
    ModelViewPermissionRequiredMixin,
    SearchMixin,
)
from apps.registry.models import Operator
from .models import KanbanBoard, KanbanBoardAccess, KanbanChecklistItem, KanbanLabel, KanbanStage, KanbanTask
from .forms import KanbanBoardForm, KanbanChecklistItemForm, KanbanLabelForm, KanbanStageForm, KanbanTaskForm


class WList(CsvExportMixin, SearchMixin, ModelViewPermissionRequiredMixin, ListView):
    template_name = "generic/list.html"
    context_object_name = "objects"
    paginate_by = 25

    def get_context_data(self, **kwargs):
        c = super().get_context_data(**kwargs)
        c["title"] = _(self.model._meta.verbose_name_plural.title())
        return c


class WCreate(ModelPermissionRequiredMixin, CreateView):
    permission_action = "add"
    template_name = "generic/form.html"
    success_url_name = None

    def get_success_url(self):
        return reverse(
            self.success_url_name or f"{self.model._meta.model_name}-list"
        )

    def form_valid(self, form):
        response = super().form_valid(form)
        if self.model is KanbanBoard and not self.object.stages.exists():
            templates = [
                (_("Pending"), "pending", "#64748B"),
                (_("Planned"), "planned", "#3B82F6"),
                (_("In progress"), "in_progress", "#F59E0B"),
                (_("Blocked"), "blocked", "#EF4444"),
                (_("In review"), "review", "#8B5CF6"),
                (_("Completed"), "completed", "#10B981"),
            ]
            KanbanStage.objects.bulk_create(
                [KanbanStage(board=self.object, name=name, status_type=status, color=color, order=index)
                 for index, (name, status, color) in enumerate(templates)]
            )
        return response

    def get_context_data(self, **kwargs):
        c = super().get_context_data(**kwargs)
        c["title"] = _("New %(record)s") % {
            "record": _(self.model._meta.verbose_name.title())
        }
        return c


class LabelList(WList):
    model = KanbanLabel


class LabelCreate(WCreate):
    model = KanbanLabel
    form_class = KanbanLabelForm
    success_url_name = "label-list"


class KanbanTaskListView(ModelViewPermissionRequiredMixin, ListView):
    model = KanbanTask
    template_name = "workboard/task_list.html"
    context_object_name = "tasks"
    permission_action = "view"
    paginate_by = 50

    def get_queryset(self):
        qs = KanbanTask.objects.filter(is_active=True, board__is_active=True).select_related("board", "stage", "assigned_to").prefetch_related("labels", "checklist_items")
        params = self.request.GET
        if params.get("board"):
            qs = qs.filter(board_id=params["board"])
        if params.get("operator"):
            qs = qs.filter(assigned_to_id=params["operator"])
        if params.get("priority") in dict(KanbanTask.PRIORITIES):
            qs = qs.filter(priority=params["priority"])
        if params.get("state") in dict(KanbanStage.STATUS_TYPES):
            qs = qs.filter(stage__status_type=params["state"])
        if params.get("label"):
            try:
                UUID(str(params["label"]))
                qs = qs.filter(labels__id=params["label"])
            except (ValueError, TypeError):
                pass
        if params.get("q"):
            qs = qs.filter(title__icontains=params["q"])
        ordering = params.get("sort")
        return qs.order_by({"priority": "priority", "due": "due_date", "assignee": "assigned_to__full_name", "progress": "updated_at"}.get(ordering, "stage__order"), "order", "created_at")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update({
            "boards": KanbanBoard.objects.filter(is_active=True),
            "operators": Operator.objects.filter(is_active=True).order_by("full_name"),
            "priorities": KanbanTask.PRIORITIES,
            "states": KanbanStage.STATUS_TYPES,
            "labels": KanbanLabel.objects.filter(board_id=self.request.GET.get("board"), is_active=True) if self.request.GET.get("board") else KanbanLabel.objects.filter(is_active=True),
            "filter_params": self.request.GET,
        })
        return context


class TaskReportCsvView(ModelViewPermissionRequiredMixin, View):
    model = KanbanTask
    permission_action = "view"

    def get(self, request):
        response = HttpResponse(content_type="text/csv; charset=utf-8")
        response["Content-Disposition"] = 'attachment; filename="aerocontrol-tasks.csv"'
        response.write("\ufeff")
        writer = csv.writer(response, lineterminator="\r\n")
        writer.writerow(["Task", "Board", "State", "Labels", "Assignee", "Priority", "Due", "Progress"])
        listing = KanbanTaskListView()
        listing.request = request
        tasks = listing.get_queryset()
        for task in tasks:
            writer.writerow([
                task.title,
                task.board.name,
                task.stage.get_status_type_display(),
                ", ".join(label.name for label in task.labels.all()),
                task.assigned_to.full_name if task.assigned_to else "",
                task.get_priority_display(),
                task.due_date.isoformat() if task.due_date else "",
                f"{task.checklist_progress}%" if task.checklist_total else "No steps",
            ])
        return response


class TaskReportXlsxView(TaskReportCsvView):
    def get(self, request):
        from io import BytesIO
        from openpyxl import Workbook
        from openpyxl.styles import Font

        listing = KanbanTaskListView()
        listing.request = request
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Tasks"
        headers = ["Task", "Board", "State", "Labels", "Assignee", "Priority", "Due", "Progress"]
        sheet.append(headers)
        for cell in sheet[1]:
            cell.font = Font(bold=True)
        for task in listing.get_queryset():
            sheet.append([
                task.title, task.board.name, task.stage.get_status_type_display(),
                ", ".join(label.name for label in task.labels.all()),
                task.assigned_to.full_name if task.assigned_to else "",
                task.get_priority_display(), task.due_date.isoformat() if task.due_date else "",
                f"{task.checklist_progress}%" if task.checklist_total else "No steps",
            ])
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for column in sheet.columns:
            sheet.column_dimensions[column[0].column_letter].width = min(max(len(str(cell.value or "")) for cell in column) + 2, 40)
        output = BytesIO()
        workbook.save(output)
        response = HttpResponse(output.getvalue(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response["Content-Disposition"] = 'attachment; filename="aerocontrol-tasks.xlsx"'
        return response


class TaskReportDocxView(TaskReportCsvView):
    def get(self, request):
        from io import BytesIO
        from docx import Document
        from docx.shared import Inches
        from django.utils import timezone

        listing = KanbanTaskListView()
        listing.request = request
        document = Document()
        document.add_heading(_("AeroControl — Operational task report"), 0)
        document.add_paragraph(_("Generated: %(date)s") % {"date": timezone.localdate().isoformat()})
        table = document.add_table(rows=1, cols=6)
        table.style = "Light Shading Accent 1"
        headers = [_("Task"), _("Board"), _("State"), _("Assignee"), _("Priority"), _("Progress")]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
        for task in listing.get_queryset():
            row = table.add_row().cells
            values = [
                task.title, task.board.name, task.stage.get_status_type_display(),
                task.assigned_to.full_name if task.assigned_to else "",
                task.get_priority_display(),
                f"{task.checklist_progress}%" if task.checklist_total else _("No steps"),
            ]
            for cell, value in zip(row, values):
                cell.text = str(value)
        for section in document.sections:
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        output = BytesIO()
        document.save(output)
        response = HttpResponse(output.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = 'attachment; filename="aerocontrol-tasks.docx"'
        return response


class ApiPermissionMixin(ModelPermissionRequiredMixin):
    def handle_no_permission(self):
        if not self.request.user.is_authenticated:
            return JsonResponse({"detail": "Authentication required."}, status=401)
        return JsonResponse({"detail": "Permission denied."}, status=403)


class ApiIndexView(LoginRequiredMixin, View):
    def get(self, request):
        return JsonResponse({
            "version": "v1",
            "authentication": "Django session or configured API gateway",
            "endpoints": {
                "tasks_list": {
                    "method": "GET",
                    "path": "/api/v1/workboard/tasks/",
                    "permission": "workboard.view_kanbantask",
                    "filters": ["board", "operator", "priority", "state", "label", "q", "page", "page_size"],
                },
                "task_update": {
                    "method": "PATCH",
                    "path": "/api/v1/workboard/tasks/<uuid>/",
                    "permission": "workboard.change_kanbantask",
                    "fields": ["title", "description", "priority", "stage_id", "due_date"],
                },
            },
        })


class ApiTaskListView(ApiPermissionMixin, View):
    model = KanbanTask
    permission_action = "view"

    def get(self, request):
        listing = KanbanTaskListView()
        listing.request = request
        queryset = listing.get_queryset()
        if not request.user.is_superuser:
            rules = KanbanBoardAccess.objects.filter(user=request.user, is_active=True)
            if rules.exists():
                queryset = queryset.filter(board_id__in=rules.values("board_id"))
        try:
            page = max(int(request.GET.get("page", "1")), 1)
            page_size = min(max(int(request.GET.get("page_size", "25")), 1), 100)
        except ValueError:
            return JsonResponse({"detail": "Invalid pagination."}, status=400)
        total = queryset.count()
        start = (page - 1) * page_size
        items = []
        for task in queryset[start:start + page_size]:
            items.append({
                "id": str(task.pk),
                "title": task.title,
                "board": {"id": str(task.board_id), "name": task.board.name},
                "state": task.stage.status_type,
                "stage": task.stage.name,
                "priority": task.priority,
                "assignee": task.assigned_to.full_name if task.assigned_to else None,
                "due_date": task.due_date.isoformat() if task.due_date else None,
                "labels": [{"id": str(label.pk), "name": label.name, "color": label.color} for label in task.labels.all()],
                "checklist": {"total": task.checklist_total, "completed": task.checklist_completed, "progress": task.checklist_progress},
                "updated_at": task.updated_at.isoformat(),
            })
        return JsonResponse({"version": "v1", "page": page, "page_size": page_size, "total": total, "results": items})


class ApiTaskUpdateView(ApiPermissionMixin, View):
    model = KanbanTask
    permission_action = "change"

    def patch(self, request, pk):
        task = get_object_or_404(KanbanTask, pk=pk, is_active=True, board__is_active=True)
        access = KanbanBoardAccess.objects.filter(board=task.board, user=request.user, is_active=True).first()
        if not request.user.is_superuser and access and access.role not in {"editor", "manager"}:
            return JsonResponse({"detail": "Object permission denied."}, status=403)
        expected_updated = request.headers.get("If-Unmodified-Since")
        if expected_updated:
            from django.utils.dateparse import parse_datetime
            expected = parse_datetime(expected_updated)
            if expected is None or task.updated_at > expected:
                return JsonResponse({"detail": "Task changed since it was read.", "code": "conflict"}, status=409)
        try:
            payload = json.loads(request.body or "{}")
        except json.JSONDecodeError:
            return JsonResponse({"detail": "Invalid JSON."}, status=400)
        allowed = {"title", "description", "priority", "stage_id", "due_date"}
        unknown = sorted(set(payload) - allowed)
        if unknown:
            return JsonResponse({"detail": "Unsupported fields.", "fields": unknown}, status=400)
        if "priority" in payload and payload["priority"] not in dict(KanbanTask.PRIORITIES):
            return JsonResponse({"detail": "Invalid priority."}, status=400)
        if "stage_id" in payload:
            stage = KanbanStage.objects.filter(pk=payload["stage_id"], board=task.board, is_active=True).first()
            if not stage:
                return JsonResponse({"detail": "Stage does not belong to this board."}, status=400)
            task.stage = stage
        for field in ("title", "description", "priority", "due_date"):
            if field in payload:
                setattr(task, field, payload[field])
        task.save()
        return JsonResponse({"version": "v1", "id": str(task.pk), "updated": sorted(payload), "updated_at": task.updated_at.isoformat()})


class TaskDetailView(ModelViewPermissionRequiredMixin, View):
    model = KanbanTask
    permission_action = "view"

    def get(self, request, pk):
        task = get_object_or_404(KanbanTask.objects.prefetch_related("labels", "checklist_items"), pk=pk, is_active=True, board__is_active=True)
        return render(request, "workboard/_task_detail.html", {"task": task, "form": KanbanTaskForm(instance=task), "checklist_form": KanbanChecklistItemForm()})


class TaskEditView(ModelPermissionRequiredMixin, UpdateView):
    model = KanbanTask
    form_class = KanbanTaskForm
    template_name = "workboard/_task_form.html"
    permission_action = "change"

    def get_queryset(self):
        return super().get_queryset().filter(is_active=True, board__is_active=True)

    def form_valid(self, form):
        response = super().form_valid(form)
        form.save_m2m()
        if self.request.headers.get("HX-Request") == "true":
            return HttpResponse(status=204, headers={"HX-Trigger": "board-refresh,task-saved"})
        return response

    def get_success_url(self):
        return reverse("kanban") + f"?board={self.object.board_id}"


class ChecklistItemCreate(ModelPermissionRequiredMixin, View):
    model = KanbanChecklistItem
    permission_action = "add"

    def post(self, request, pk):
        task = get_object_or_404(KanbanTask, pk=pk, is_active=True, board__is_active=True)
        form = KanbanChecklistItemForm(request.POST)
        if form.is_valid():
            item = form.save(commit=False)
            item.task = task
            item.order = task.checklist_items.count()
            item.save()
        return render(request, "workboard/_task_detail.html", {"task": task, "form": KanbanTaskForm(instance=task), "checklist_form": KanbanChecklistItemForm()})


class ChecklistItemToggle(ModelPermissionRequiredMixin, View):
    model = KanbanChecklistItem
    permission_action = "change"

    def post(self, request, pk):
        item = get_object_or_404(KanbanChecklistItem, pk=pk, task__is_active=True)
        item.is_completed = not item.is_completed
        item.save(update_fields=["is_completed", "updated_at"])
        task = item.task
        return render(request, "workboard/_task_detail.html", {"task": task, "form": KanbanTaskForm(instance=task), "checklist_form": KanbanChecklistItemForm()})


for model, form, name in (
    (KanbanBoard, KanbanBoardForm, "Board"),
    (KanbanStage, KanbanStageForm, "Stage"),
    (KanbanTask, KanbanTaskForm, "Task"),
):
    globals()[f"{name}List"] = type(f"{name}List", (WList,), {"model": model})
    globals()[f"{name}Create"] = type(
        f"{name}Create",
        (WCreate,),
        {
            "model": model,
            "form_class": form,
            "success_url_name": f"{name.lower()}-list",
        },
    )


# ─── Kanban Board Views ─────────────────────────────────────


class KanbanBoardView(LoginRequiredMixin, TemplateView):
    """Main kanban board view — all stages and task cards."""

    template_name = "workboard/kanban.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        boards = KanbanBoard.objects.filter(is_active=True)
        board_id = self.request.GET.get("board")
        board = self._get_board(board_id)
        context.update(
            {
                "board": board,
                "boards": boards,
                "operators": self._get_operators(),
                "priorities": KanbanTask.PRIORITIES,
                "states": KanbanStage.STATUS_TYPES,
                "labels": KanbanLabel.objects.filter(board=board, is_active=True) if board else KanbanLabel.objects.none(),
                "drag_enabled": self._filter_values(self.request.GET)[0] is None
                and not self._filter_values(self.request.GET)[1],
                "stages": self._build_stage_data(board, self.request.GET)
                if board
                else [],
                "filter_params": self.request.GET,
            }
        )
        return context

    @staticmethod
    def _filter_values(params):
        operator_id = params.get("operator", "")
        priority = params.get("priority", "")
        try:
            operator = (
                Operator.objects.filter(pk=operator_id, is_active=True).first()
                if operator_id
                else None
            )
        except (ValueError, TypeError, ValidationError):
            operator = None
        if priority not in dict(KanbanTask.PRIORITIES):
            priority = ""
        return operator, priority

    @staticmethod
    def _get_board(board_id):
        boards = KanbanBoard.objects.filter(is_active=True)
        if not board_id:
            return boards.first()
        try:
            return boards.filter(pk=board_id).first() or boards.first()
        except (ValueError, TypeError, ValidationError):
            return boards.first()

    def _build_stage_data(self, board, params):
        stages = board.stages.filter(is_active=True).order_by("order")
        operator, priority = self._filter_values(params)
        state = params.get("state") if params.get("state") in dict(KanbanStage.STATUS_TYPES) else ""
        label = params.get("label")
        try:
            if label:
                UUID(str(label))
        except (ValueError, TypeError):
            label = ""
        query = params.get("q", "").strip()
        stage_data = []
        for stage in stages:
            tasks = stage.tasks.filter(is_active=True).order_by("order")
            if operator:
                tasks = tasks.filter(assigned_to=operator)
            if priority:
                tasks = tasks.filter(priority=priority)
            if state:
                tasks = tasks.filter(stage__status_type=state)
            if label:
                tasks = tasks.filter(labels__id=label)
            if query:
                tasks = tasks.filter(title__icontains=query)
            stage_data.append({"stage": stage, "tasks": tasks})
        return stage_data

    def _get_operators(self):
        return Operator.objects.filter(is_active=True).order_by("full_name")


class BoardPartialView(LoginRequiredMixin, View):
    """HTMX fragment — board columns + cards for filter/drag refresh."""

    def get(self, request):
        if request.headers.get("HX-Request") != "true":
            query = request.GET.urlencode()
            target = reverse("kanban")
            return redirect(f"{target}?{query}" if query else target)
        board_id = request.GET.get("board")
        board = KanbanBoardView._get_board(board_id)

        if not board:
            return HttpResponse(
                '<p class="text-muted text-center py-5">No board configured.</p>'
            )

        stage_data = KanbanBoardView()._build_stage_data(board, request.GET)
        response = render(
            request,
            "workboard/_board.html",
            {
                "board": board,
                "stages": stage_data,
                "drag_enabled": KanbanBoardView._filter_values(request.GET)[0] is None
                and not KanbanBoardView._filter_values(request.GET)[1],
                "filter_params": request.GET,
            },
        )
        query = request.GET.urlencode()
        response.headers["HX-Push-Url"] = (
            f"{reverse('kanban')}?{query}" if query else reverse("kanban")
        )
        return response


class StageCreate(ModelPermissionRequiredMixin, CreateView):
    model = KanbanStage
    form_class = KanbanStageForm
    permission_action = "add"
    template_name = "generic/form.html"

    def get_initial(self):
        initial = super().get_initial()
        if self.request.GET.get("board"):
            initial["board"] = self.request.GET["board"]
        return initial

    def get_success_url(self):
        board_id = self.object.board_id
        return f"{reverse('kanban')}?board={board_id}"


class BoardArchiveView(ModelPermissionRequiredMixin, View):
    model = KanbanBoard
    permission_action = "change"

    def post(self, request, pk):
        board = get_object_or_404(KanbanBoard, pk=pk, is_active=True)
        board.is_active = False
        board.save(update_fields=["is_active", "updated_at"])
        return redirect("kanban")


class TaskArchiveView(ModelPermissionRequiredMixin, View):
    model = KanbanTask
    permission_action = "change"

    def post(self, request, pk):
        task = get_object_or_404(
            KanbanTask, pk=pk, is_active=True, board__is_active=True
        )
        board_id = task.board_id
        task.is_active = False
        task.save(update_fields=["is_active", "updated_at"])
        return redirect(f"{reverse('kanban')}?board={board_id}")


class MoveTaskView(ModelPermissionRequiredMixin, View):
    """Drag-and-drop: move task between stages and re-order siblings."""

    model = KanbanTask
    permission_action = "change"

    def post(self, request, pk):
        task = get_object_or_404(
            KanbanTask,
            pk=pk,
            is_active=True,
            board__is_active=True,
            stage__is_active=True,
        )
        old_stage_id = task.stage_id

        stage_id = request.POST.get("stage_id")
        raw_order = request.POST.get("new_order", "0")

        try:
            new_order = max(int(raw_order), 0)
            new_stage = KanbanStage.objects.filter(
                pk=stage_id,
                board_id=task.board_id,
                is_active=True,
            ).first()
        except (ValueError, TypeError, ValidationError):
            return HttpResponse(status=400)
        if new_stage is None or task.stage.board_id != task.board_id:
            return HttpResponse(status=400)

        with transaction.atomic():
            old_stage = KanbanStage.objects.get(pk=old_stage_id)
            self._remove_from_stage(task, old_stage)
            task.stage = new_stage
            task.save(update_fields=["stage", "updated_at"])
            self._insert_into_stage(task, new_stage, new_order)

        return HttpResponse(status=204, headers={"HX-Trigger": "board-refresh"})

    @staticmethod
    def _remove_from_stage(task, stage):
        siblings = list(
            stage.tasks.filter(is_active=True)
            .exclude(pk=task.pk)
            .order_by("order", "created_at")
        )
        for index, sibling in enumerate(siblings):
            if sibling.order != index:
                KanbanTask.objects.filter(pk=sibling.pk).update(order=index)

    @staticmethod
    def _insert_into_stage(task, stage, order):
        siblings = list(
            stage.tasks.filter(is_active=True)
            .exclude(pk=task.pk)
            .order_by("order", "created_at")
        )
        siblings.insert(min(order, len(siblings)), task)
        for index, sibling in enumerate(siblings):
            if sibling.pk == task.pk:
                if task.order != index:
                    KanbanTask.objects.filter(pk=task.pk).update(order=index)
            elif sibling.order != index:
                KanbanTask.objects.filter(pk=sibling.pk).update(order=index)


class QuickTaskCreate(ModelPermissionRequiredMixin, View):
    """Quick-add a task from the column footer, returns refreshed column."""

    model = KanbanTask
    permission_action = "add"

    def get(self, request):
        """Return the quick-add form fragment for a given stage."""
        stage_id = request.GET.get("stage_id")
        if not stage_id:
            return HttpResponse(status=400)
        try:
            stage = KanbanStage.objects.filter(
                pk=stage_id,
                is_active=True,
                board__is_active=True,
            ).first()
        except (ValueError, TypeError, ValidationError):
            stage = None
        if stage is None:
            return HttpResponse(status=400)
        operators = KanbanBoardView()._get_operators()
        return render(
            request,
            "workboard/_quick_form.html",
            {
                "stage": stage,
                "operators": operators,
                "filter_params": request.GET,
            },
        )

    def post(self, request):
        stage_id = request.POST.get("stage_id")
        title = request.POST.get("title", "").strip()
        assigned_to = request.POST.get("assigned_to") or None
        priority = request.POST.get("priority", "medium")

        if not title or not stage_id or priority not in dict(KanbanTask.PRIORITIES):
            return HttpResponse(status=400)

        try:
            stage = KanbanStage.objects.filter(
                pk=stage_id,
                is_active=True,
                board__is_active=True,
            ).first()
        except (ValueError, TypeError, ValidationError):
            stage = None
        if stage is None:
            return HttpResponse(status=400)
        if assigned_to:
            try:
                operator = Operator.objects.filter(
                    pk=assigned_to, is_active=True
                ).first()
            except (ValueError, TypeError, ValidationError):
                return HttpResponse(status=400)
            if operator is None:
                return HttpResponse(status=400)
        max_order = KanbanTask.objects.filter(stage=stage, is_active=True).count()

        KanbanTask.objects.create(
            board=stage.board,
            stage=stage,
            title=title,
            priority=priority,
            assigned_to_id=assigned_to,
            created_by=request.user.get_username(),
            order=max_order,
        )

        filter_params = request.POST.copy()
        filter_priority = filter_params.get("filter_priority", "")
        if filter_priority:
            filter_params["priority"] = filter_priority
        else:
            filter_params.pop("priority", None)
        operator, valid_priority = KanbanBoardView._filter_values(filter_params)
        tasks = stage.tasks.filter(is_active=True).order_by("order")
        if operator:
            tasks = tasks.filter(assigned_to=operator)
        if valid_priority:
            tasks = tasks.filter(priority=valid_priority)
        return render(
            request,
            "workboard/_column.html",
            {
                "stage": stage,
                "tasks": tasks,
                "drag_enabled": operator is None and not valid_priority,
                "filter_params": filter_params,
            },
        )


class BoardSelector(LoginRequiredMixin, View):
    """Board switcher dropdown fragment."""

    def get(self, request):
        boards = KanbanBoard.objects.filter(is_active=True)
        current = request.GET.get("board")
        return render(
            request,
            "workboard/_board_selector.html",
            {
                "boards": boards,
                "current": current,
            },
        )


# Backwards-compatible name used by earlier workboard integrations.
BoardFilterPartial = BoardPartialView
